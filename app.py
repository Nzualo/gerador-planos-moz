import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
import pandas as pd
from docx import Document
import io
import time

# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(page_title="Gestão de Planos", page_icon="🇲🇿", layout="wide")

# --- FUNÇÃO DE LOGIN ---
def check_password():
    if st.session_state.get("password_correct", False):
        return True

    st.markdown("## 🇲🇿 SNE - Elaboração de Planos de Aulas")
    st.markdown("##### Serviço Distrital de Educação, Juventude e Tecnologia - Inhassoro")
    st.divider()
    
    col1, col2 = st.columns([1, 1])
    with col1:
        st.info("🔐 Área Restrita")
        usuario = st.text_input("Usuário")
        senha = st.text_input("Senha", type="password")
        if st.button("Entrar", type="primary"):
            if "passwords" in st.secrets and usuario in st.secrets["passwords"]:
                if st.secrets["passwords"][usuario] == senha:
                    st.session_state["password_correct"] = True
                    st.session_state["user_name"] = usuario
                    st.rerun()
                else:
                    st.error("Senha incorreta.")
            else:
                st.error("Usuário desconhecido.")

    with col2:
        st.warning("⚠️ Suporte")
        st.write("Precisa de acesso? Fale com o Administrador.")
        meu_numero = "258867926665"
        mensagem = "Olá Técnico Nzualo, gostaria de solicitar acesso ao Sistema de Planos."
        link_zap = f"https://wa.me/{meu_numero}?text={mensagem.replace(' ', '%20')}"
        st.markdown(f'<a href="{link_zap}" target="_blank"><button style="background-color:#25D366; color:white; border:none; padding:10px 20px; border-radius:5px; width:100%; cursor:pointer;">📱 Contactar via WhatsApp</button></a>', unsafe_allow_html=True)
    return False

if not check_password():
    st.stop()

# --- BARRA LATERAL ---
with st.sidebar:
    st.success(f"👤 Olá, **{st.session_state['user_name']}**")
    if st.button("Sair"):
        st.session_state["password_correct"] = False
        st.rerun()

# --- FUNÇÕES DE DOCUMENTOS ---
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 5, 'REPÚBLICA DE MOÇAMBIQUE', 0, 1, 'C')
        self.set_font('Arial', 'B', 10)
        self.cell(0, 5, 'GOVERNO DO DISTRITO DE INHASSORO', 0, 1, 'C')
        self.cell(0, 5, 'SDEJT - INHASSORO', 0, 1, 'C')
        self.ln(5)
        self.set_font('Arial', 'B', 14)
        self.cell(0, 10, 'PLANO DE AULA', 0, 1, 'C')
        self.ln(2)
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 6)
        self.cell(0, 10, 'SDEJT Inhassoro - Processado por IA', 0, 0, 'C')
    def table_row(self, data, widths):
        max_lines = 1
        for i, text in enumerate(data):
            self.set_font("Arial", size=8)
            lines = self.multi_cell(widths[i], 4, text, split_only=True)
            if len(lines) > max_lines: max_lines = len(lines)
        height = max_lines * 4 + 4
        if self.get_y() + height > 270:
            self.add_page()
            self.ln(5)
        x_start = self.get_x()
        y_start = self.get_y()
        for i, text in enumerate(data):
            self.set_xy(x_start, y_start)
            self.multi_cell(widths[i], 4, text, border=1)
            self.set_xy(x_start + widths[i], y_start)
            x_start += widths[i]
        self.set_y(y_start + height)

def create_pdf(inputs, dados, objetivos):
    pdf = PDF()
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 7, f"Escola: ___________________________________  Data: ___/___/____", 0, 1)
    pdf.cell(0, 7, f"Disciplina: {inputs['disciplina']} | Classe: {inputs['classe']} | Turma: {inputs['turma']}", 0, 1)
    pdf.cell(0, 7, f"Unidade: {inputs['unidade']}", 0, 1)
    pdf.set_font("Arial", "B", 10)
    pdf.cell(0, 7, f"Tema: {inputs['tema']}", 0, 1)
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 7, f"Tipo: {inputs['tipo_aula']} | Duração: {inputs['duracao']}", 0, 1)
    pdf.ln(3)
    pdf.set_font("Arial", "B", 10)
    pdf.multi_cell(0, 5, f"OBJETIVOS: {objetivos}")
    pdf.ln(3)
    widths = [12, 28, 35, 35, 35, 22, 23]
    headers = ["Tempo", "F. Didática", "Conteúdo", "Professor", "Aluno", "Métodos", "Meios"]
    pdf.set_font("Arial", "B", 8)
    pdf.set_fill_color(220, 220, 220)
    for i, h in enumerate(headers):
        pdf.cell(widths[i], 6, h, 1, 0, 'C', True)
    pdf.ln()
    for row in dados:
        pdf.table_row(row, widths)
    return pdf.output(dest='S').encode('latin-1', 'ignore')

def create_word(inputs, dados, objetivos):
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run('REPÚBLICA DE MOÇAMBIQUE\nGOVERNO DO DISTRITO DE INHASSORO\nSDEJT\n\nPLANO DE AULA')
    run.bold = True
    doc.add_paragraph(f"Escola: __________________ Data: ____/____/____")
    doc.add_paragraph(f"Disciplina: {inputs['disciplina']} | Classe: {inputs['classe']}")
    doc.add_paragraph(f"Tema: {inputs['tema']}")
    doc.add_paragraph(f"Objetivos: {objetivos}")
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ["Tempo", "Função", "Conteúdo", "Prof", "Aluno", "Métodos", "Meios"]
    for i, h in enumerate(headers): hdr_cells[i].text = h
    for row_data in dados:
        row_cells = table.add_row().cells
        for i, item in enumerate(row_data): row_cells[i].text = str(item)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel(dados):
    df = pd.DataFrame(dados, columns=["Tempo", "Função", "Conteúdo", "Prof", "Aluno", "Métodos", "Meios"])
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    buffer.seek(0)
    return buffer

# --- INTERFACE PRINCIPAL ---
st.title("🇲🇿 Elaboração de Planos de Aulas")

if "GOOGLE_API_KEY" not in st.secrets:
    st.error("⚠️ Erro: Configure os Secrets.")
    st.stop()

col1, col2 = st.columns(2)
with col1:
    disciplina = st.text_input("Disciplina", "Língua Portuguesa")
    classe = st.selectbox("Classe", ["1ª", "2ª", "3ª", "4ª", "5ª", "6ª", "7ª", "8ª", "9ª", "10ª", "11ª", "12ª"])
    unidade = st.text_input("Unidade", placeholder="Ex: Textos Normativos")
    tipo_aula = st.selectbox("Tipo", ["Inicial", "Exercitação", "Revisão", "Avaliação"])
with col2:
    duracao = st.selectbox("Duração", ["45 Min", "90 Min"])
    turma = st.text_input("Turma", placeholder="A")
    tema = st.text_input("Tema", placeholder="Ex: Vogais")

# --- BOTÃO GERAR ---
if st.button("🚀 Criar Plano (PDF, Word, Excel)", type="primary"):
    with st.spinner('A processar...'):
        try:
            genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
            model = genai.GenerativeModel('models/gemini-2.5-flash')
            prompt = f"""
            Aja como Pedagogo do SNE Moçambique.
            Plano: {disciplina}, {classe}, Tema: {tema}.
            REGRAS: 1. TPC (Correção/Marcação). 2. OBJETIVOS: Max 3. 3. TABELA: "||".
            SAÍDA: [BLOCO_OBJETIVOS]...[FIM_OBJETIVOS] [BLOCO_TABELA]...[FIM_TABELA]
            """
            response = model.generate_content(prompt)
            texto = response.text
            
            # Extração de Dados
            objetivos = "..."
            dados = []
            if "[BLOCO_OBJETIVOS]" in texto:
                objetivos = texto.split("[BLOCO_OBJETIVOS]")[1].split("[FIM_OBJETIVOS]")[0].strip()
            if "[BLOCO_TABELA]" in texto:
                lines = texto.split("[BLOCO_TABELA]")[1].split("[FIM_TABELA]")[0].strip().split('\n')
                for l in lines:
                    if "||" in l and "Função" not in l:
                        cols = [c.strip() for c in l.split("||")]
                        while len(cols) < 7: cols.append("-")
                        dados.append(cols)
            
            # GUARDAR NA MEMÓRIA (O SEGREDO!)
            st.session_state['plano_gerado'] = True
            st.session_state['dados'] = dados
            st.session_state['objetivos'] = objetivos
            st.session_state['inputs'] = {'disciplina': disciplina, 'classe': classe, 'duracao': duracao, 'tema': tema, 'unidade': unidade, 'tipo_aula': tipo_aula, 'turma': turma}
            st.rerun() # Recarrega a página para mostrar os botões
            
        except Exception as e:
            st.error(f"Erro: {e}")

# --- MOSTRAR RESULTADOS (SÓ SE JÁ TIVER GERADO) ---
if st.session_state.get('plano_gerado'):
    st.divider()
    st.subheader("✅ Plano Gerado com Sucesso!")
    
    # Recuperar dados da memória
    dados = st.session_state['dados']
    objetivos = st.session_state['objetivos']
    inputs = st.session_state['inputs']

    st.info(f"**Objetivos:** {objetivos}")
    if dados:
        df = pd.DataFrame(dados, columns=["Tempo", "Função", "Conteúdo", "Prof", "Aluno", "Métodos", "Meios"])
        st.dataframe(df, hide_index=True)

    st.markdown("### 📥 Baixar Documentos")
    c1, c2, c3 = st.columns(3)
    
    # Gerar arquivos na hora do clique
    pdf_bytes = create_pdf(inputs, dados, objetivos)
    c1.download_button("📄 PDF (Oficial)", data=pdf_bytes, file_name=f"Plano_{inputs['disciplina']}.pdf", mime="application/pdf")
    
    word_bytes = create_word(inputs, dados, objetivos)
    c2.download_button("📝 Word (Editável)", data=word_bytes, file_name=f"Plano_{inputs['disciplina']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    excel_bytes = create_excel(dados)
    c3.download_button("📊 Excel (Tabela)", data=excel_bytes, file_name=f"Plano_{inputs['disciplina']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    
    if st.button("🔄 Criar Novo Plano"):
        st.session_state['plano_gerado'] = False
        st.rerun()
